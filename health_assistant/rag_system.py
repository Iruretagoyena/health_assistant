from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from health_assistant.config import set_openai_key
from langchain.schema import Document
from docx import Document as DocxDocument
import pickle
import faiss
import os

api_key = set_openai_key()
embedding_model = OpenAIEmbeddings(api_key=api_key)

def parse_docx_to_texts(file_path):
    doc = DocxDocument(file_path)
    texts = []
    for para in doc.paragraphs:
        # Collect non-empty paragraphs as separate "documents"
        if para.text.strip():
            texts.append(para.text.strip())
    return texts

def create_faiss_index(texts, embedding_model):
    # Convert texts to Document objects
    documents = [Document(page_content=text) for text in texts]
    
    # Create FAISS index from documents
    faiss_index = FAISS.from_documents(documents, embedding_model)
    
    # Save the FAISS index itself using faiss.write_index
    faiss.write_index(faiss_index.index, "faiss_index.idx")
    
    # Save the docstore (or any additional components) with pickle
    with open("faiss_docstore.pkl", "wb") as f:
        pickle.dump(faiss_index.docstore, f)  # Save only the docstore separately
    print("FAISS index and docstore saved successfully.")
    return faiss_index

def create_rag_system(llm):
    # Load or create the FAISS index
    try:
        faiss_index_obj = faiss.read_index("faiss_index.idx")
        with open("faiss_docstore.pkl", "rb") as f:
            docstore = pickle.load(f)  # Load the docstore directly
        
        # Create a new index_to_docstore_id mapping
        index_to_docstore_id = {i: id_ for i, id_ in enumerate(docstore._dict.keys())}
        
        faiss_index = FAISS(
            embedding_function=embedding_model,
            index=faiss_index_obj,
            docstore=docstore,
            index_to_docstore_id=index_to_docstore_id
        )
        print("FAISS index loaded successfully.")
    except RuntimeError:
        print("FAISS index not found. Creating a new index...")
        # texts = [
        #     "Sleep hygiene practices can improve sleep quality.",
        #     "Regular exercise helps maintain a healthy heart rate.",
        #     "Nutritional foods contribute to overall well-being."
        # ]
        print(os.getcwd())
        texts = parse_docx_to_texts("documents/sleep_knowledge.docx")
        # print(texts)
        faiss_index = create_faiss_index(texts, embedding_model)
        print("FAISS index created and saved.")

    # Set up the retriever from the FAISS index
    retriever = faiss_index.as_retriever()

    # Create and return the RetrievalQA chain
    return RetrievalQA.from_chain_type(
        llm=llm,
        retriever=retriever
    )